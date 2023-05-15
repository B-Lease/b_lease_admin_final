from flask import make_response
from fpdf import FPDF
import os
from datetime import datetime, date
from docx import Document
from io import BytesIO
from PIL import Image, UnidentifiedImageError

from docx.shared import Inches
#from 
import util
import docx
import docx2pdf
import pythoncom
import json
import base64

unchecked = u'\u2b1c '
checked = u'\u2611 '
def setContract(leasingInfo, contractInfo):
    date_obj = datetime.strptime(str(date.today()), '%Y-%m-%d').strftime('%B %d, %Y')
    leasing_start = datetime.strptime(leasingInfo['leasing_start'], '%Y-%m-%d').strftime('%B %d, %Y')
    leasing_end = datetime.strptime(leasingInfo['leasing_end'], '%Y-%m-%d').strftime('%B %d, %Y')

    day = ordinal_suffix(int(datetime.strptime(leasingInfo['leasing_start'], '%Y-%m-%d').strftime('%d')))
    month_day = day + datetime.strptime(leasingInfo['leasing_start'], '%Y-%m-%d').strftime(' of %B')

    leasing_payment_frequency = leasingInfo['leasing_payment_frequency']
    one_time = checked if leasing_payment_frequency == '1' else unchecked
    monthly = checked if leasing_payment_frequency == '2' else unchecked
    monthly_date = day if leasing_payment_frequency == '2' else '___'
    
    quarterly = checked if leasing_payment_frequency == '3' else unchecked
    quarterly_date = day if leasing_payment_frequency == '3' else '___'

    annually = checked if leasing_payment_frequency == '4' else unchecked
    annually_date = month_day if leasing_payment_frequency == '4' else '___'

    if contractInfo['security_deposit']:
        security_deposit = checked
        non_security_deposit = unchecked
    else:
        security_deposit = unchecked
        non_security_deposit = checked

    if contractInfo['improvements']:
        improvements = checked
        non_improvements = unchecked
    else:
        improvements = unchecked
        non_improvements = checked

    if contractInfo['erect_signage']:
        erect_signage = checked
        non_erect_signage = unchecked
    else:
        erect_signage = unchecked
        non_erect_signage = checked

    if leasingInfo['leasing_payment_type'] == 'cash rent':
        fee = float(leasingInfo['leasing_total_fee'])
        leasing_total_fee = "Php "+ "{:,.2f}".format(fee)
    else:
        leasing_total_fee = leasingInfo['leasing_total_fee'] + "%" + " of the crops produced during the total lease term"

    signature = contractInfo['signature']

    # Remove the prefix using replace()
    # signature = signature.replace('data:image/png;base64,', '')

    #for filename
    leasingID = leasingInfo['leasingID']
    target_input_map = {
        'This Land/Ground Lease Agreement (this “Agreement”) is entered into as of ': date_obj,
        'Lessor: ': contractInfo['lessor_name'], 
        'Lessee: ': contractInfo['lessee_name'], 
        'according to the terms and conditions set forth herein, the following real estate (the “Site”): ': contractInfo['land_description'],
        'The Site may be used and occupied only for the following purpose (the “Permitted Use”): ': contractInfo['purpose'],
        'This Agreement will be for a term ': 'beginning on ' + leasing_start + ' and ending on ' + leasing_end,
        'Lessee will pay Lessor rent in the form of ': leasingInfo['leasing_payment_type'],
        'totaling to ':leasing_total_fee,
        'On the Start Date': one_time,
        'In monthly': monthly,
        'In monthly installments due on the': ' ' + monthly_date,
        'Due monthly lease rental paid after the ': monthly_date,
        'In annual installments ': annually,
        'Due annual lease rental paid after ': annually_date,
        'In quarterly installments': quarterly,
        'In quarterly installments due on the': ' ' + quarterly_date,
        'Due quarterly lease rental paid after ': quarterly_date,
        'Lessee is NOT required to pay a security deposit.': non_security_deposit,
        'Upon signing this Agreement, Lessee will pay a security deposit': security_deposit,
        'may not may make improvements': non_improvements,
        'may make improvements, alterations': improvements,
        'Lessee shall have the right to erect any sign ': erect_signage,
        'Lessee shall NOT have the right to erect any sign': non_erect_signage,
        'IN WITNESS WHEREOF, the Parties hereto, individually or by their duly authorized representatives have executed this Agreement as of the Effective Date.': signature
    }

    #print(json.dumps(target_input_map, indent=4))

    doc_path = "static\\contracts\\lease-agreement.docx"
    doc = Document(doc_path)

    for target_string, input_string in target_input_map.items():
        for p in doc.paragraphs:
            if target_string in p.text:
                # Split the paragraph into two parts: before and after the target string
                parts = p.text.split(target_string)
                
                # Set font properties for the entire new paragraph
                p.style.font.name = 'Arial'
                p.style.font.size = docx.shared.Pt(10)
                
                if input_string == unchecked or input_string == checked:
                    p.text = f"{parts[0]}"
                    p.add_run(input_string)
                    p.add_run(f"{target_string}{parts[1]}")
                
                elif input_string == signature:
                    
                    p.text = f"{parts[0]}{target_string}"

                    # Decode the base64 image string
                    image_data = BytesIO(base64.b64decode(signature))

                    # Add the image to the document
                    p.add_run('\n\nLessor signature: \n').bold = True
                    p.add_run().add_picture(image_data, width=Inches(3))
                    p.add_run('\n\nLessee signature: \n').bold = True
                    p.add_run(parts[1])

                else:
                    # Insert first character until the target string
                    # Insert the new string.upper() and set it to bold
                    # Insert the proceeding string after the new string
                    p.text = f"{parts[0]}{target_string}"
                    p.add_run(input_string.upper()).bold = True
                    p.add_run(parts[1])
                    break

    # Specify the path of the parent directory
    parent_dir = "static\\contracts\\"

    # Specify the name of the new directory to create
    new_dir = f"{leasingID}"

    # Create the full path of the new directory
    path = os.path.join(parent_dir, new_dir)

    # Create the new directory
    make_path = print('ok') if os.path.exists(path) else os.makedirs(path)

    docx_path = f"static\\contracts\\{leasingID}\\{util.generateUUID(str(datetime.now()))}_pending.docx" 
    
    doc.save(f"{docx_path}")

    # pdf_path = f"static\\contracts\\{leasingID}\\{util.generateUUID(str(datetime.now()))}_ongoing.pdf"
    # # Initialize COM
    # pythoncom.CoInitialize()

    # # Convert to PDF
    # docx2pdf.convert(docx_path, pdf_path)

    # # Uninitialize COM
    # pythoncom.CoUninitialize()

    # os.remove(docx_path)
    return True

def signContract(contractInfo):
    leasingID = contractInfo['leasingID']
    signature = contractInfo['signature']

    doc_path = f"static\\contracts\\{leasingID}\\"

    # Get a list of all files in the folder
    files = os.listdir(doc_path)

    # Get the first file in the folder
    first_file = files[0]

    # Print the filename
    # Open the Word document
    doc = docx.Document(doc_path + first_file)


    # Get the last paragraph in the document
    last_paragraph = doc.paragraphs[-1]

    # Add a new paragraph to the end of the document
    new_paragraph = doc.add_paragraph()

    # Add your text to the new paragraph
    image_data = BytesIO(base64.b64decode(signature))
    new_paragraph.add_run().add_picture(image_data, width=Inches(3))

    
    doc.save(doc_path + first_file)

    pdf_path = f"static\\contracts\\{leasingID}\\{util.generateUUID(str(datetime.now()))}_for_review.pdf"
    # Initialize COM
    pythoncom.CoInitialize()

    # Convert to PDF
    docx2pdf.convert(doc_path + first_file, pdf_path)

    # Uninitialize COM
    pythoncom.CoUninitialize()

    os.remove(doc_path + first_file)
    return contractInfo

def ordinal_suffix(day):
    if day in [11, 12, 13]:
        suffix = "th"
    elif day % 10 == 1:
        suffix = "st"
    elif day % 10 == 2:
        suffix = "nd"
    elif day % 10 == 3:
        suffix = "rd"
    else:
        suffix = "th"
    
    return str(day) + suffix
    
#==============================================================================================================
#==============================================================================================================
#==============================================================================================================
#==============================================================================================================

