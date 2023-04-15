from flask import make_response
from fpdf import FPDF
import os
from datetime import datetime, date
from docx import Document
from io import BytesIO
from PIL import Image, UnidentifiedImageError

from docx.shared import Inches
#from 
import util
import docx
import docx2pdf
import pythoncom
import json
import base64

unchecked = u'\u2b1c '
checked = u'\u2611 '
def createPDF(leasingInfo, contractInfo):
    date_obj = datetime.strptime(str(date.today()), '%Y-%m-%d').strftime('%B %d, %Y')
    leasing_start = datetime.strptime(leasingInfo['leasing_start'], '%Y-%m-%d').strftime('%B %d, %Y')
    leasing_end = datetime.strptime(leasingInfo['leasing_end'], '%Y-%m-%d').strftime('%B %d, %Y')

    day = ordinal_suffix(int(datetime.strptime(leasingInfo['leasing_start'], '%Y-%m-%d').strftime('%d')))
    month_day = day + datetime.strptime(leasingInfo['leasing_start'], '%Y-%m-%d').strftime(' of %B')

    leasing_payment_frequency = leasingInfo['leasing_payment_frequency']
    one_time = checked if leasing_payment_frequency == '1' else unchecked
    monthly = checked if leasing_payment_frequency == '2' else unchecked
    monthly_date = day if leasing_payment_frequency == '2' else '___'
    
    quarterly = checked if leasing_payment_frequency == '3' else unchecked
    quarterly_date = day if leasing_payment_frequency == '3' else '___'

    annually = checked if leasing_payment_frequency == '4' else unchecked
    annually_date = month_day if leasing_payment_frequency == '4' else '___'

    if contractInfo['security_deposit']:
        security_deposit = checked
        non_security_deposit = unchecked
    else:
        security_deposit = unchecked
        non_security_deposit = checked

    if contractInfo['improvements']:
        improvements = checked
        non_improvements = unchecked
    else:
        improvements = unchecked
        non_improvements = checked

    if contractInfo['erect_signage']:
        erect_signage = checked
        non_erect_signage = unchecked
    else:
        erect_signage = unchecked
        non_erect_signage = checked


    signature = contractInfo['signature']

    # Remove the prefix using replace()
    # signature = signature.replace('data:image/png;base64,', '')

    #for filename
    leasingID = leasingInfo['leasingID']
    target_input_map = {
        'This Land/Ground Lease Agreement (this “Agreement”) is entered into as of ': date_obj,
        'Lessor: ': contractInfo['lessor_name'], 
        'Lessee: ': contractInfo['lessee_name'], 
        'according to the terms and conditions set forth herein, the following real estate (the “Site”): ': contractInfo['land_description'],
        'The Site may be used and occupied only for the following purpose (the “Permitted Use”): ': contractInfo['purpose'],
        'This Agreement will be for a term ': 'beginning on ' + leasing_start + ' and ending on ' + leasing_end,
        'Lessee will pay Lessor rent in total of Php ':leasingInfo['leasing_total_fee'],
        'On the Start Date': one_time,
        'In monthly': monthly,
        'In monthly installments due on the': ' ' + monthly_date,
        'Due monthly lease rental paid after the ': monthly_date,
        'In annual installments ': annually,
        'Due annual lease rental paid after ': annually_date,
        'In quarterly installments': quarterly,
        'In quarterly installments due on the': ' ' + quarterly_date,
        'Lessee is NOT required to pay a security deposit.': non_security_deposit,
        'Upon signing this Agreement, Lessee will pay a security deposit': security_deposit,
        'may not may make improvements': non_improvements,
        'may make improvements, alterations': improvements,
        'Lessee shall have the right to erect any sign ': erect_signage,
        'Lessee shall NOT have the right to erect any sign': non_erect_signage,
        'IN WITNESS WHEREOF, the Parties hereto, individually or by their duly authorized representatives have executed this Agreement as of the Effective Date.': signature
    }

    #print(json.dumps(target_input_map, indent=4))

    doc_path = "static\\contracts\\lease-agreement.docx"
    doc = Document(doc_path)

    for target_string, input_string in target_input_map.items():
        for p in doc.paragraphs:
            if target_string in p.text:
                # Split the paragraph into two parts: before and after the target string
                parts = p.text.split(target_string)
                
                # Set font properties for the entire new paragraph
                p.style.font.name = 'Arial'
                p.style.font.size = docx.shared.Pt(10)
                
                if input_string == unchecked or input_string == checked:
                    p.text = f"{parts[0]}"
                    p.add_run(input_string)
                    p.add_run(f"{target_string}{parts[1]}")
                
                elif input_string == signature:
                    
                    p.text = f"{parts[0]}{target_string}"

                    # Decode the base64 image string
                    image_data = BytesIO(base64.b64decode(signature))

                    # Add the image to the document
                    p.add_run('\n\nLessor signature: \n').bold = True
                    p.add_run().add_picture(image_data, width=Inches(3))
                    p.add_run(parts[1])

                else:
                    # Insert first character until the target string
                    # Insert the new string.upper() and set it to bold
                    # Insert the proceeding string after the new string
                    p.text = f"{parts[0]}{target_string}"
                    p.add_run(input_string.upper()).bold = True
                    p.add_run(parts[1])
                    break

    docx_path = "static\\contracts\\land-lease-agreement.docx" 
    
    doc.save(f"{docx_path}")
    
    # Specify the path of the parent directory
    parent_dir = "static\\contracts\\"

    # Specify the name of the new directory to create
    new_dir = f"{leasingID}"

    # Create the full path of the new directory
    path = os.path.join(parent_dir, new_dir)

    # Create the new directory
    make_path = print('ok') if os.path.exists(path) else os.makedirs(path)

    pdf_path = f"static\\contracts\\{leasingID}\\{util.generateUUID(str(datetime.now()))}_ongoing.pdf"
    # Initialize COM
    pythoncom.CoInitialize()

    # Convert to PDF
    docx2pdf.convert(docx_path, pdf_path)

    # Uninitialize COM
    pythoncom.CoUninitialize()

    os.remove(docx_path)
    return True

def ordinal_suffix(day):
    if day in [11, 12, 13]:
        suffix = "th"
    elif day % 10 == 1:
        suffix = "st"
    elif day % 10 == 2:
        suffix = "nd"
    elif day % 10 == 3:
        suffix = "rd"
    else:
        suffix = "th"
    
    return str(day) + suffix
    
#==============================================================================================================
#==============================================================================================================
#==============================================================================================================
#==============================================================================================================

from flask import Flask, make_response
from fpdf import FPDF
import re

#portrait layout, mm unit of measurement, and letter format (short bp)
pdf = FPDF('P', 'mm', 'Letter')

def generate_pdf():
    #VARIABLES
    state = 'Cebu'.upper()
    date = '12'.upper()
    month = 'January'.upper()
    year = '2023'.upper()
    lessor = "Allain Zenith N. Sabandal".upper()
    lessee = "Bianca N. Sabandal".upper()

    # Create a new PDF object
    pdf = FPDF()
    # Add a page to the PDF
    pdf.add_page()

    #set font
    pdf.set_font("Arial")
    
    #write text
    pdf.set_font("", size=8.5)
    text = f"State of [{state}]"
    writePDF(pdf, text)
    
    #write text
    text = f"Rev. 133EF48"
    pdf.cell(0,5,text,align='R')
    pdf.ln()

    #set font
    pdf.set_font('Arial', 'B', 16.5)

    #title
    pdf.cell(0, 10, 'LAND/GROUND LEASE AGREEMENT', align='C')
    pdf.ln()
    pdf.line(10, pdf.get_y(), pdf.w - 10, pdf.get_y())
    pdf.ln()

    #set font
    pdf.set_font('Arial', '', 10)
    #write text
    text = (f"""This Land/Ground Lease Agreement (this "Agreement") is entered into as of the [{date}th] day of [{month}], [{year}], """
        f"""(the "Effective Date") by and among/between: \n\n"""
        f"""Lessor: [{lessor}], ('Lessor') and \n"""
        f"""Lessee(s): [{lessee}], (collectively, 'Lessee').\n\n"""

        f"""[1. Agreement to Lease.] Lessor agrees to lease to Lessee and Lessee agrees to lease from Lessor, according to the """
        f"""terms and conditions set forth herein, the following real estate (the "Site"): Lots 6, 7, and the South ½ of Lot 3, West """
        f"""60 feet of South ½ of Lot 4, West 60 feet of Lot 5 and Lot 8, Block 20, OLD SURVEY, Leesville, Vernon Parish, Louisiana. """
        f"""[Legal land description].\n\n"""

        f"""[2. Purpose.] The Site may be used and occupied only for the following purpose (the "Permitted Use"): Commercial use [Purpose of land use]. """
        f"""Nothing herein shall give Lessee the right to use the Site for any other purpose without the prior written consent of Lessor."""
        f"""Lessor makes no representation or warranty regarding the legality of the Permitted Use, and Lessee will bear all risk of any adverse"""
        f"""change in applicable laws.\n\n"""
        
        f"""[3. Term.] This Agreement will be for a term beginning on [January 19, 2023] and ending on [January 19, 2024] (the "Term")."""
        f"""The Parties hereto may elect to extend this Agreement upon such terms and conditions as may be agreed upon in writing and signed by"""
        f"""the Parties at the time of any such extension.\n\n"""

        f"""[4. Rent.] Lessee will pay Lessor rent in advance Php_____a_____: (Check one) \n"""
        
        )


    writePDF(pdf,text)
    
    pdf.set_left_margin(15)
    add_checkmark(pdf,True)
    writePDF(pdf, '\t On the Effective Date. \n')

    add_checkmark(pdf,False)
    writePDF(pdf, '\t In monthly installments due on the _____a_____ day of each month during the Term. \n')

    pdf.set_left_margin(20.5)
    text = (
        f"""[Late Fee]:\n"""
        f"""Rent paid after the _____a_____ day of each month will be deemed as late; and if rent is not paid within """
        f"""_____a_____days after such due date, Lessee agrees to pay a late charge of _____a_____% of the balance due """
        f"""per day for eachday that rent is late."""  
        )
    
    writePDF(pdf, text)

    pdf.set_left_margin(15)
    pdf.write(5,f"""\n""")
    
    add_checkmark(pdf, False)
    writePDF(pdf, f"""\tIn annual installments due on the _____a_____ [Day of the month] day of _______a____ [Month] each year \
        \t\t\t\t\tduring the Term.\n""")
    
    pdf.set_left_margin(20.5)

    text = (
        f"""[Late Fee]:\n"""
        f"""Rent paid after the _____a_____ [Day of the month] day of _______a________ [Month] of each year will be deemed """
        f"""as late; and if rent is not paid within _____a_____ days after such due date, Lessee agrees to pay a late charge of   """
        f"""_____a_____% of the balance due per day for each day that rent is late."""  
    )
    
    writePDF(pdf, text)

    pdf.set_left_margin(15)
    pdf.write(5,f"""\n""")

    add_checkmark(pdf,False)
    writePDF(pdf, '\t In quarterly installments due on the _____a_____ day of the ending month of each quarter during the Term.')

    pdf.set_left_margin(10)
    pdf.write(5,f"""\n""")
    text = (
        f"""[5. Additional Rent.] There may be instances under this Agreement where Lessee may be required to pay additional charges to Lessor. """
        f"""All such charges are considered additional rent under this Agreement and will be paid with the next regularly scheduled rent payment."""
        f"""Lessor has the same rights and Lessee has the same obligations with respect to additional rent as they do with rent. \n\n\n"""
        f"""[6. Security Deposit.]\n"""
    )

    writePDF(pdf, text)

    add_checkmark(pdf,True)
    writePDF(pdf, '\t Lessee is NOT required to pay a security deposit. \n')

    add_checkmark(pdf,False)
    text = (f"""\t Upon signing this Agreement, Lessee will pay a security deposit in the amount of $_____a_____ to Lessor. The security """
        f"""deposit will be retained by Lessor as security for Lessee's performance of its obligations under this Agreement. If Lessee """
        f"""does not comply with any of the terms of this Agreement, Lessor may apply any or all of the security deposit to remedy the """
        f"""breach, including to cover any amount owed by Lessee and/or any damages or costs incurred by Lessor due to Lessee's failure """
        f"""to comply. Within _____a_____ days after the termination of this Agreement, Lessor will return the security deposit to Lessee """
        f"""(minus any amount applied by Lessor in accordance with this section). Any reason for retaining a portion of the security deposit """
        f"""will be explained in writing.\n\n"""
        
        f"""\t\t\t[Interest]:\n\t\t\t"""
        f"""The security deposit will not bear interest.\n\n"""
        
        f"""[7. Taxes.] Lessee shall pay all taxes or assessments which are levied or charged on the Site during the Term. The lease rentals """
        f"""collected by the Lessor shall include the payments for the said taxes or assessments.\n\n"""
        
        f"""[8. Utilities.] Lessee shall pay the cost of all utility services during the Term, including but not limited to gas, water, and """
        f"""electricity used on the Site.\n\n"""

        f"""[9. Delivery of Possession.] Landlord will deliver exclusive and lawful possession of the Site to Tenant on the """
        f"""start date of the Term. In the event Landlord is unable to give possession of the Site to Tenant on such date, """
        f"""Landlord will not be subject to any liability for such failure, the validity of this Agreement will not be """
        f"""affected, and the Term will not be extended. Tenant will not be liable for rent until Landlord gives possession """
        f"""of the Site to Tenant.\n\n"""
        
        f"""[10. Holdover Tenancy.] Unless this Agreement has been extended by mutual written agreement of the Parties, """
        f"""there will be no holding over past the Term under the terms of this Agreement under any circumstances. If """
        f"""Tenant does retain possession past the Term, Tenant shall pay ____a______% of the then applicable rent """
        f"""computed on a monthly basis for each month or portion thereof during such holdover. In addition, Tenant """
        f"""shall be liable for any damages incurred by Landlord as a result of the holdover.\n\n"""
        
        f"""[11. Condition of the Site.] Tenant has examined the Site and accepts the Site in its current condition "AS IS\" """
        f"""and "WITH ALL FAULTS." except as expressly set forth herein, landlord makes no representation OR warranty, """
        f"""express or implied, or arising by operation of law, including but not limited to, any warranty of fitness for """
        f"""a particular purpose, merchantability, habitability, SUITABILITY, or condition. tenant acknowledges that Tenant """
        f"""has not relied on any representations or warranties by Landlord in entering this Agreement. \n\n"""
        
        f"""[12. Use of the Site.] Tenant agrees to use the Site only for the Permitted Use and will not commit waste """
        f"""upon the Site. Tenant will, at its sole expense, maintain the Site in good repair and make all necessary """
        f"""repairs thereto. Tenant will not use the Site for any unlawful purpose or in any manner that will materially """
        f"""harm Landlord's interest in the Site.\n\n [13. Improvements and Alterations.] """

    )

    writePDF(pdf, text)
    
    pdf.write(5, '\n\t\tLessee \t\t\t')
    add_checkmark(pdf,False)
    pdf.write(5, ' may not \t\t\t')
    add_checkmark(pdf,True)
    
    text = (f""" may make improvements, alterations, additions, or other changes to the Site without the """
        f"""\t\twritten approval of the Landlord. Tenant agrees that any construction will be performed in a good and workmanlike """
        f"""\t\tmanner and will comply with all applicable laws. All improvements, alterations, additions, or other changes to the Site """
        f"""\t\tshall become the property of Landlord upon the termination of this Agreement."""
    )

    writePDF(pdf,text)

    pdf.write(5, '\n\t\tSigns:\n\t\t\t')
    add_checkmark(pdf,False)
    pdf.write(5, ' Tenant shall have the right to erect any sign related to its business, on the condition that such signs comply with \t\t\t\t\t\t the law.\n\t\t\t')
    add_checkmark(pdf,True)
    
    text = (f""" Tenant shall NOT have the right to erect any sign related to its business. \n"""
        f"""[14. Leasehold Mortgage.] """   
    )
    
    writePDF(pdf, text)

    pdf.write(5, '\n\t\tLessee \t\t\t')
    add_checkmark(pdf,False)
    pdf.write(5, ' does not have \t\t\t')
    add_checkmark(pdf,True)
    
    text = (f""" has the right to grant a mortgage, deed of trust, or other security instrument in Tenant's """
        f"""interest to the Site created by this Agreement (the "Leasehold Mortgage") to secure repayment of a loan made """
        f"""to Tenant to finance construction of any improvements made to the Site during the Term. """
    )

    writePDF(pdf,text)

    # Subordinated Lease (Check one)  a
    # ☐ Landlord agrees to pledge as collateral or subordinate its interest in the Site for or to any Leasehold Mortgage if required by any lender of Tenant. 
    # ☐ In no event will any interest of Landlord in the Site be pledged as collateral for or be subordinate to any Leasehold Mortgage.


    #=====================================================
    
    # Create a response object with the PDF data
    #response = make_response(pdf.output(dest='S').encode('utf-16'))
    response = make_response(pdf.output(dest='S').encode('latin1'))
    
    # Set the content type of the response to PDF
    response.headers.set('Content-Type', 'application/pdf')
    
    # Return the response object
    return response


def writePDF(pdf, text:str):
    # # Define the regular expression pattern to match interpolated expressions
    # pattern = r"\[\s*(\w+(?:\s+\w+)*)\s*\]"

    # # Split the string into parts based on the pattern
    # parts = re.split(pattern, text)

    pattern = r"\[\s*([^\]]+)\s*\]"
    parts = re.split(pattern, text)
    print(parts)
    print(len(parts))
    for i in range(len(parts)):
        if (i % 2 == 0):
            # The part is regular text, set the font color to black
            pdf.set_font("", "")
            pdf.write(5,parts[i])
        else:
            # The part is an interpolated expression, set the font color to green
            pdf.set_font("", "B")
            pdf.write(5,parts[i])

    pdf.ln()

def add_checkmark(pdf,checked:bool):
    # Set font to ZapfDingbats
    pdf.set_font('ZapfDingbats', '', 10)
    if checked:
    # Add check mark symbol
        pdf.cell(4,4, chr(52), border=1)
    else:
        pdf.cell(4,4, ' ', border=1)
    pdf.set_font('Arial','',10)

